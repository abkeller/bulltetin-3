from docx import Document
#from docx import WE_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.shared import Pt
import os
from pathlib import Path
import docx

import sys
sys.path.append("..")

#from .utils import pct, strint
from utils.doc_utils import *
from utils.garages import GARAGES, GAR_MANAGER, WEEKDAYS, DAY_TYPES

from datetime import datetime as dt

GRAY = '#E7E7E7'

#bulletin = {
#        'key': 'value'
#        }

export_dir = 'M:\\akeller\\Python\\Bulletin Creation\\Bulletins-2.0\\Bulletins\\python\\templates'

class CubsBulletin:
    def __init__(self):
        self.intro = [
                """ add a line of text broken each time it's r: regular, b: bold, u: underlined
                code will determine how to display text based on first two characters of line [0:2]
                code will print beginning with the third character [3:]
                place spaces at end of text line to maintain proper spacing in document
                """
                
                'r: To accommodate additional ridership for a ',
                'b: Chicago Cubs ',
                'r: game at ',
                'b: Wrigley Field ',
                'r: beginning at',
                'u: 1320 ',
                'r: on ',
                'u: {} '.format('Sundays'),
                'r: operate ',
                'r: {} on route {} '.format('two (2) extras', '#80-Irving Park'),
                'r: and ',
                'r: {} on route {} '.format('nine (9) extras', '#152-Addison'),
                'r: on ',
                'r: five (5) dates ',
                'r: during the ',
                'u: Summer 2019 pick',
                'r: .'
                ]
    
        self.bullet1 = [
                'r: Gates at ',
                'b: Wrigley Field',
                'r: will open at ',
                'u: 1120',
                'r:.'
                ]
        
        self.bullet2 = [
                'r: The game is scheduled to begin at ',
                'u: 1320',
                'r: and end at apprximately ',
                'u: 1550',
                'r:.'
                ]
        
        self.bullets = [self.bullet1, self.bullet2]
        self.box_data = {
                'Number': 'SB19-0283',
                'Effective': 'June 23, 2019',
                'Copies to': 'Bus Distribution List',
                'Supersedes': '',
                'Activity code': '11078'
                }
        
        self.day_type = 'su'
        self.dates = ['JUNE 23', 'JULY 14', 'JULY 21', 'AUGUST 4', 'AUGUST 25']
        self.re_line = 'RE: #80-Irving Park/ #152-Addison - Chicago Cubs Sunday games @ 1320 - Summer 2019 pick'
        self.dates1 = {
                'headers': ['MONDAY', 'TUESDAY', 'WEDNESDAY', 'THURSDAY', 'FRIDAY'],
                1: {
                        1: '',
                        2: '',
                        3: '',
                        4: '04/01/2021',
                        5: ''
                        },
                2: {
                        1: '',
                        2: '',
                        3: '4/7/2021',
                        4: '',
                        5: ''
                        }  
        }

class BulletinCover:
    def __init__(self, bulletin, export_dir):
        self.bulletin = bulletin
        self.export_dir = export_dir

        # Create document and add styles defined in doc_utils
        self.doc = Document()
        edit_styles(self.doc)
        set_margins(self.doc, Inches(0.5))

        # add document headers
        headers = ["CHICAGO TRANSIT AUTHORITY", "Scheduling & Service Planning", "Bus Schedule Bulletin"]
        self.page_headers(headers)
        
        ## add text box        
        self.text_box(self.bulletin)
       
        ## add garage headers, send date
        garage = 'P'
        send_date = 'August 5, 2019'
        self.add_garage_headers(garage, send_date)

        ## add intro paragraph
        self.intro_paragraph(self.bulletin)
        
        ## add bullets
        self.add_bullets(self.bulletin)
        
        # add re line
        self.add_re_line(self.bulletin)

        # add schedule
        self.add_schedule(self.bulletin)      

        ##########################################################################
        doc_name = 'test.docx'
        doc_path = os.path.join('M:\\akeller\\Python\\Bulletin Creation\\Bulletins-2.0\\Bulletins\\python\\templates', doc_name)
        self.doc.save(doc_path) 
        ############################################################################



    def add_schedule(self, bulletin):
        rows = len(bulletin.dates1)
        cols = len(bulletin.dates1['headers'])
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = getattr(WD_TABLE_ALIGNMENT,'CENTER')
        
        for i, d in enumerate(bulletin.dates1):
            #print(i)
            if i == 0:
                for i, h in enumerate(bulletin.dates1['headers']):
                    cell = table.cell(0, i)
                    cell.width = Inches(2)
                    cell.text = h
                    #print(h)
            elif i > 0:
                for j, d in enumerate(bulletin.dates1[i]):
                    print(j)
                    print(bulletin.dates1[i][d])
                    if bulletin.dates1[i][d] != '':
                        date = dt.strptime(bulletin.dates1[i][d], '%m/%d/%Y')
                        date = date.strftime('%B %#d')
                        print(date)
                        cell = table.columns[j].cells[i]
                        cell.text = date
#        #table.allow_autofit = True
#        #table(autofit=True)
#        table.cell(0, 0).text = "REAlldy long text wheat will happen"
#        #table.cell(0, 0).text = DAY_TYPES[bulletin.day_type]['name'].upper()
#        thick_black = {'sz': 24, 'val': 'single', 'color': '#000000'}
#
#        for i, row in enumerate(table.rows):
#            cell = table.columns[0].cells[i]
#            if i > 0:
#                cell.text = bulletin.dates[i - 1]
#                cell.width = table.allow_autofit
#            set_cell_border(cell, start=thick_black, end=thick_black, top=thick_black, bottom=thick_black)
#            
#            
#            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
#            row.height = Pt(14)
            
        #table.allow_autofit = True
        


    def add_re_line(self, bulletin):
        para = self.doc.add_paragraph(bulletin.re_line)
        para.style = self.doc.styles['UL']
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


    def add_bullets(self, bulletin):

        for b in bulletin.bullets:
            para = self.doc.add_paragraph("")
            para.style = self.doc.styles['List Bullet 2']
            para.height = Pt(12)
        
            for text in b:
                if text[0:2] == 'r:':
                    para.add_run(text[3:])
    
                elif text[0:2] == 'b:':
                    para.add_run(text[3:]).bold = True             
    
                elif text[0:2] == 'u:':
                    para.add_run(text[3:]).underline = True 
        

    def intro_paragraph(self, bulletin):
        para = self.doc.add_paragraph("\n")
        para.style = self.doc.styles['Paragraph']
        

        for text in bulletin.intro:
            if text[0:2] == 'r:':
                para.add_run(text[3:])

            elif text[0:2] == 'b:':
                para.add_run(text[3:]).bold = True             

            elif text[0:2] == 'u:':
                para.add_run(text[3:]).underline = True           


    def add_garage_headers(self, garage, send_date):
        # create as table so manager can be left justified, and send date right justified
        table = self.doc.add_table(rows=2, cols=2)
        table.style = self.doc.styles['GH']
        
        # add manager headers
        manager_cell = table.cell(0, 0)
        manager_cell.text = '{}, {}'.format(GAR_MANAGER[garage]['sm_name'], GAR_MANAGER[garage]['sm_title'])
        
        # add date on right side
        date_cell = table.cell(0, 1)
        date_cell.text = send_date
        format_cells(date_cell, halign='right', valign='bottom')
        
        # add garage name on second row
        garage_cell = table.cell(1, 0)
        garage_cell.text = '{} Garage'.format(GARAGES[garage])
        
        for i, row in enumerate(table.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Pt(14)

    # creates the headers at the top of the bulletin page
    def page_headers(self, headers):
                # create 
#        styles = self.doc.styles
        
        for h in headers:
            header = self.doc.add_paragraph(h)
            header.style = self.doc.styles['H1']



    # creates the text box at top left corner        
    def text_box(self, bulletin):
        data = bulletin.box_data
        length = len(data)
        table = self.doc.add_table(rows=length, cols=2)
        table.alignment = getattr(WD_TABLE_ALIGNMENT,'RIGHT')
        table.style = self.doc.styles['TB']
        set_col_widths(table, (Inches(1.2), Inches(2.0)))

        for i, row in enumerate(table.rows):
            #print(row)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            if i == 0:
                row.height = Pt(20)
            else:
                row.height = Pt(16)

        for i, val in enumerate(data):
            table.cell(i, 0).text = val
            table.cell(i, 1).text = data[val]
            
            
            thick_black = {'sz': 24, 'val': 'single', 'color': '#000000'}
            for cell in table.rows[i].cells:
                format_cells(cell, bold=True, size=Pt(12), valign='bottom')
                if i == 0:
                    set_cell_border(cell, top=thick_black)
                if i == length - 1:
                    set_cell_border(cell, bottom=thick_black)
            
                set_cell_border(table.columns[0].cells[i], start=thick_black)
                set_cell_border(table.columns[1].cells[i], end=thick_black)
                
#        word = win32com.client.Dispatch('Word.Appplication')













        
        
        
bulletin = CubsBulletin()
BulletinCover(bulletin, export_dir)

#        # Write report to PDF
#        word = win32com.client.Dispatch('Word.Application')
##        try:
##            docx = word.Documents.Open(doc_path)
##            try:
#        pdf_name = 'test.pdf'
#        pdf_path = os.path.join(doc_path, pdf_name)
#        docx.SaveAs(pdf_path, FileFormat=17)
#
#
#            doc = word.Documents.Open(path)
#            doc.SaveAs(path.replace('.docx', '.pdf'), FileFormat=17)
#            doc.Close()#            except:
##                print('      Error printing PDF report.')
##            docx.Close()
##        except:
###            print('      Error printing PDF report.')
#
#        word.Quit()

    # ************************************************************************
            

#
#        # Add Calculation notes for Tripper Breakdown - TO-DO
#        if bd_type == 'tripper_bd':
#            calc_notes = ['','','Calculation = Total tripper pieces per week divided by 5',
#                          'Weekday trippers are multiplied by 5 because there are 5 weekdays per week',
#                          'Saturday and Sunday runs set aside for PTOs are multiplied by 2 because there are 2 pieces per run']
#
#            for cn in calc_notes:
#                para = self.doc.add_paragraph(cn)
#                para.style = self.doc.styles['Footnote']
#
#        return table
#
#    # ************************************************************************
#    # Pages 7-9: Interline Current Pick (Wk/Sa/Su)
#    def interline_table(self, day_type):
#        self.add_headers('interline_table', day_type=day_type)
#        curr_data = self.run_char['curr']['interline'][day_type]
#
#        # Create table and set headers
#        cols = 4 if day_type == 'wk' else 3
#        table = self.doc.add_table(rows=4 * len(curr_data) + 6, cols=cols)
#        if day_type == 'wk':
#            set_col_widths(table, (11.16, Inches(2.1), Inches(0.75),
#                                   Inches(0.75)))
#        else:
#            set_col_widths(table, (Inches(0.9), Inches(2.1), Inches(1.25)))
#        table.style = 'Table Grid'
#
#       # create borders     
#        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
#        blank = {'val': 'none', 'color': '#ffffff'}
#        row_length = len(table.rows)
#        for i in range(row_length):
#           for cell in table.rows[i].cells:
#               set_cell_border(cell, start=thick_black, end=thick_black)
#               
#           if i == 0:
#               for cell in table.rows[i].cells:
#                   set_cell_border(cell, top=thick_black, bottom=blank)
#                   set_cell_border(table.columns[i].cells[i], start=blank, top=blank)
#                   
#           if i == 1:
#               for cell in table.rows[i].cells:
#                   set_cell_border(cell, top=blank, bottom=thick_black)
#                   set_cell_border(table.columns[0].cells[i], start=blank)
#                  
#           if (i - 5) % 4 == 0 and i >= 5:
#               for cell in table.rows[i].cells:
#                   set_cell_border(cell, top=thick_black, bottom=thick_black)
#                   set_cell_border(table.columns[0].cells[i], start=blank)
#                   set_cell_border(table.columns[cols - 1].cells[i])
#
#           if i == 28:
#               for cell in table.rows[i].cells:
#                   set_cell_border(cell, bottom=thick_black)
#                    
#           if i > 28:
#               for cell in table.rows[i].cells:
#                   set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)
#
#        # add column headers to table            
#        headers = ('\nGarage', '\nRun Type',
#                   '\nRegular Runs', '\nBlock')        
#        for i, c in enumerate(headers):
#            if day_type in ['su', 'sa'] and i == 4:
#                pass
#            elif day_type == 'su' and i == 2:
#                c = '\nReg. FTO'
#                table.cell(0, i).text = c
#                format_cells(table.cell(0, i), bold=True)
#            else:
#                table.cell(0, i).text = c
#                format_cells(table.cell(0, i), bold=True)
#
#        tr_keys, tr_type = get_keys('interline')
#        run_types = ['reg', 'block'] if day_type == 'wk' else ['reg']
#        
#        for i, g in enumerate(curr_data):
#            # First column - Garage (merged)
#            gar = table.cell(i * 4 + 2, 0).merge(table.cell(i * 4 + 4, 0))
#            gar.text = GARAGES.get(g)
#            format_cells(gar, bold=True)
#
#            for j in range(3):
#                # Second column - Run Type
#                table.cell(i * 4 + 2 + j, 1).text = tr_type[j]
#
#                # Third-Fifth columns - Regular Runs
#                # Sixth-Eighth columns (Weekday) - Block Runs
#                is_pct = tr_keys[j] == 'pct'
#                row = i * 4 + 2 + j
#                
#                for k, run_type in enumerate(run_types):
#                    curr_ct = curr_data[g][run_type][tr_keys[j]]
#                    table.cell(row, 2 + k).text = strint(curr_ct, pct=is_pct)
#
#        # System Total
#        sys_start = len(curr_data) * 4 + 3
#        system = table.cell(sys_start, 0).merge(table.cell(sys_start + 2, 0))
#        system.text = 'System'
#        format_cells(system, bold=True)
#
#        for j in range(3):
#            table.cell(sys_start + j, 1).text = tr_type[j]
#
#        for i, run_type in enumerate(run_types):
#            iline, total = [], []
#            for j, data in enumerate(curr_data):
#                #print(curr_data[data])
#                iline.append(int(curr_data[data][run_type]['iline']))
#                total.append(int(curr_data[data][run_type]['total']))
#            iline_total = sum(iline)
#            total = sum(total)
#
#            col = i + 2
#            table.cell(sys_start, col).text = strint(iline_total)
#            table.cell(sys_start + 1, col).text = strint(total)
#            table.cell(sys_start + 2, col).text = strint(pct(iline_total, total), pct=True)
#
#        # Apply Table Cell style
#        apply_table_cell(self.doc, table)
#
#        # Draw thick borders and shade every other garage
#        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
#        for j in range(len(curr_data) + 1):
#            for cell in table.rows[j * 3 + 2].cells:
#                pass
#                #set_cell_border(cell, top=thick_black)
#            if j % 2 == 1:
#                for k in range(3):
#                    for cell in table.rows[j * 4 - k].cells:
#                        set_cell_shading(cell, GRAY)
#
#        return table
#
#
#    # ************************************************************************
#    # Pages 10-12: Interline Comparisons (Wk/Sa/Su)
#    def interline_table_comp(self, day_type):
#        self.add_headers('interline_table', day_type=day_type)
#        curr_data = self.run_char['curr']['interline'][day_type]
#        prev_data = self.run_char['prev']['interline'][day_type]
#
#        # Create table and set headers
#        cols = 8 if day_type == 'wk' else 5
#        table = self.doc.add_table(rows=4 * len(curr_data) + 7, cols=cols)             
#        
#        if day_type == 'wk':
#            set_col_widths(table, (Inches(0.9), Inches(2.1), Inches(0.75),
#                                   Inches(0.75), Inches(0.75), Inches(0.75),
#                                   Inches(0.75), Inches(0.75)))
#        else:
#            set_col_widths(table, (Inches(0.9), Inches(2.1), Inches(1.25),
#                                   Inches(1.25), Inches(1.25)))
#        table.style = 'Table Grid'
#        row_length = len(table.rows)
#       # create borders     
#        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
#        blank = {'val': 'none', 'color': '#ffffff'}
#                 
#        if day_type == 'wk':
#            for i in range(row_length):
#                if  i == 0:
#                    for cell in table.rows[i].cells:
#                        set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)
#                        set_cell_border(table.columns[i].cells[i], start=blank, bottom=blank)
#                
#                if i == 1:
#                    for cell in table.rows[i].cells:
#                        set_cell_border(table.columns[0].cells[i], start=blank, bottom=blank, top=blank)
#
#                if (i - 5) % 4 == 0 and i >= 5:
#                   for cell in table.rows[i].cells:
#                       #set_cell_border(cell, top=thick_black, bottom=thick_black)
#                       set_cell_border(table.columns[0].cells[i], start=blank)
#                       set_cell_border(table.columns[cols - 1].cells[i], end=blank)            
#
#                if i > 28:
#                   for cell in table.rows[i].cells:
#                       set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)
#
#        if day_type != 'wk':
#            for i in range(row_length):
#               for cell in table.rows[i].cells:
#                   set_cell_border(cell, start=thick_black, end=thick_black)
#                   
#               if i == 0:
#                   for cell in table.rows[i].cells:
#                       set_cell_border(cell, top=blank, bottom=thick_black, start=blank, end=blank)
#                       set_cell_border(table.columns[0].cells[i], start=blank, bottom=blank, top=blank)
#
#                       
#               if i == 1:
#                   for cell in table.rows[i].cells:
#                       set_cell_border(cell, top=blank, bottom=thick_black)
#                      
#               if (i - 5) % 4 == 0 and i >= 5:
#                   for cell in table.rows[i].cells:
#                       set_cell_border(cell, top=thick_black, bottom=thick_black)
#                       set_cell_border(table.columns[0].cells[i], start=blank)
#                       set_cell_border(table.columns[cols - 1].cells[i], end=blank)
#                       set_cell_border(table.columns[0].cells[1], start=blank, top=blank)
#
#    
#               if i == 28:
#                   for cell in table.rows[i].cells:
#                       set_cell_border(cell, bottom=thick_black)
#                        
#               if i > 28:
#                   for cell in table.rows[i].cells:
#                       set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)
#
#        reg = table.cell(0, 2)
#        reg.text = '{}\n'.format(self.prev_pick['desc'].upper().split('BUS SYSTEM - ')[1])
#        curr = table.cell(0, 3)
#        curr.text = '{}\n'.format(self.curr_pick['desc'].upper().split('BUS SYSTEM - ')[1])
#        for cell in [reg, curr]:
#            format_cells(cell, bold=True, italic=True)
#        
#        if day_type == 'wk':
#            prev = table.cell(0, 5)
#            prev.text = '{}\n'.format(self.prev_pick['desc'].upper().split('BUS SYSTEM - ')[1])
#            curr = table.cell(0, 6)
#            curr.text = '{}\n'.format(self.curr_pick['desc'].upper().split('BUS SYSTEM - ')[1]) 
#            
#            for cell in [prev, curr]:
#                format_cells(cell, bold=True, italic=True)
#            
#        for i, c in enumerate(table.rows[1].cells):
#            headers = ('Garage', 'Run Type\n\n', 'Regular Runs','Regular Runs', 'Change',
#                       'Block', 'Block', 'Change')
#
#            c.text = headers[i]
#            format_cells(c, bold=True)
#
#        tr_keys, tr_type = get_keys('interline')
#        run_types = ['reg', 'block'] if day_type == 'wk' else ['reg']
#
#        for i, g in enumerate(curr_data):
#            # First column - Garage (merged)
#            gar = table.cell(i * 4 + 2, 0).merge(table.cell(i * 4 + 4, 0))
#            gar.text = GARAGES.get(g)
#            format_cells(gar, bold=True)
#
#            for j in range(3):
#                # Second column - Run Type
#                table.cell(i * 4 + 2 + j, 1).text = tr_type[j]
#
#                # Third-Fifth columns - Regular Runs
#                # Sixth-Eighth columns (Weekday) - Block Runs
#                is_pct = tr_keys[j] == 'pct'
#                row = i * 4 + 2 + j
#                
#                for k, run_type in enumerate(run_types):
#                    prev_ct = prev_data[g][run_type][tr_keys[j]]
#                    curr_ct = curr_data[g][run_type][tr_keys[j]]
#                    table.cell(row, 2 + k * 3).text = strint(prev_ct, pct=is_pct)
#                    table.cell(row, 3 + k * 3).text = strint(curr_ct, pct=is_pct)
#                    table.cell(row, 4 + k * 3).text = strint(curr_ct - prev_ct,
#                                                             pct=is_pct)
#
#        # System Total
#        sys_start = len(curr_data) * 4 + 4
#        system = table.cell(sys_start, 0).merge(table.cell(sys_start + 2, 0))
#        system.text = 'System'
#        format_cells(system, bold=True)
#
#
#        for j in range(3):
#            table.cell(sys_start + j, 1).text = tr_type[j]
#
#        for i, run_type in enumerate(run_types):
#            iline, total = [], []
#            for j, data in enumerate((prev_data, curr_data)):
#                iline.append(sum([data[g][run_type]['iline'] for g in data]))
#                total.append(sum([data[g][run_type]['total'] for g in data]))
#
#                col = 2 + j + 3 * i
#                table.cell(sys_start, col).text = strint(iline[j])
#                table.cell(sys_start + 1, col).text = strint(total[j])
#                table.cell(sys_start + 2, col).text = strint(pct(iline[j], total[j]), pct=True)
#
#            col = 4 + 3 * i
#            table.cell(sys_start, col).text = strint(iline[1] - iline[0])
#            table.cell(sys_start + 1, col).text = strint(total[1] - total[0])
#            table.cell(sys_start + 2, col).text = strint(pct(iline[1], total[1]) -
#                                                         pct(iline[0], total[0]),
#                                                         pct=True)
#
#        # Apply Table Cell style
#        apply_table_cell(self.doc, table)
#
#        # Draw thick borders and shade every other garage
#        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
#        for j in range(len(curr_data) + 1):
#            for cell in table.rows[j * 4 + 2].cells:
#                pass
#                #set_cell_border(cell, top=thick_black)
#            if j % 2 == 1:
#                for k in range(3):
#                    for cell in table.rows[j * 4 - k].cells:
#                        set_cell_shading(cell, GRAY)
#
#        return table
#    # ************************************************************************
#
#    def add_headers(self, table_type, day_type=None):
#        start_date = self.curr_pick['start_date'].strftime('%#d %B %Y').upper()
#        h1 = self.doc.add_heading('CTA Bus System Effective {} ({})'.format(
#            start_date, self.curr_pick['name']), level=1)
#        h1.style = self.doc.styles['H1']
#
#        if table_type == 'run_change_table':
#            h2_text = 'Profile of Run Changes by Garage - {}'.format(
#                DAY_TYPES[day_type]['name'].upper())
#            h3_text = 'Composite Changes in Run Categories for {} Schedules'.format(
#                DAY_TYPES[day_type]['name'])
#        elif table_type == 'tripper_bd':
#            h2_text = 'Breakdown of trippers and work set aside for part-time operators'
#        elif table_type == 'block_pct':
#            h2_text = 'Breakdown of Block Runs'
#        elif table_type == 'swing_pct':
#            h2_text = 'Breakdown of Swing Runs'
#        elif table_type == 'interline_table':
#            h2_text = 'Breakdown of Runs Working 2 or More Routes'        
#        elif table_type == 'interline_table_comp':
#            h2_text = 'Changes in Runs Working 2 or More Routes - {}'.format(
#                DAY_TYPES[day_type]['name'].upper())
#        else:
#            h2_text = ''
#        
#        h2 = self.doc.add_heading(h2_text, level=2)
#        h2.style = self.doc.styles['H2']
#
#        if table_type == 'tripper_bd':
#            h3_text = 'Calculated total of trippers must not exceed 1500'
#            h3 = self.doc.add_paragraph(h3_text)
#            h3.style = self.doc.styles['H2']
#            h3.runs[0].font.italic = True
#
#        elif table_type == 'block_pct':
#            h3_text = 'Block Runs Must not exceed 20% of the total WK & Sa FTO Runs'
#            h3 = self.doc.add_paragraph(h3_text)
#            h3.style = self.doc.styles['H2']
#            h3.runs[0].font.italic = True
#
#        elif table_type == 'swing_pct':
#            h3_text = 'Swing Runs must not exceed 30% of the total Wk & Sat FTO runs'
#            h3 = self.doc.add_paragraph(h3_text)
#            h3.style = self.doc.styles['H2']
#            h3.runs[0].font.italic = True 
#            
#        elif table_type == 'interline_table':
#            h3_text = DAY_TYPES[day_type]['name']
#            h3 = self.doc.add_paragraph(h3_text)
#            h3.style = self.doc.styles['H2']
#            h3.runs[0].font.italic = True  
#
## Gets keys used to identify table data and their printed names
#def get_keys(bd_type):
#    if bd_type == 'tripper_bd':
#        tr_keys = ['wk_tripper', 'sa_tripper', 'sa_pto', 'su_pto']
#        tr_type = ['Weekday trippers',
#                   'Saturday trippers',
#                   'Saturday runs set aside for PTOs',
#                   'Sunday runs set aside for PTOs']
#    elif bd_type == 'block_pct':
#        tr_keys = ['reg', 'block', 'fto', 'pct']
#        tr_type = ['Non-block FTO Runs ((Wk*5)+Sa)',
#                   'Block FTO Runs ((Wk*5)+Sa)',
#                   'TOTAL FTO Runs ((Wk*5)+Sa)',
#                   'Blocks as a % of Total']
#    elif bd_type == 'swing_pct':
#        tr_keys = ['sw', 'swb', 'fto', 'pct']
#        tr_type = ['Regular Swing ((Wk*5)+Sa)',
#                   'Block Swing ((Wk*5+Sa)',
#                   'TOTAL FTO Runs ((Wk*5)+Sa)',
#                   'Swings as a % of Total']
#    elif bd_type == 'interline':
#        tr_keys = ['iline', 'total', 'pct']
#        tr_type = ['Runs working 2+ routes',
#                   'TOTAL Runs',
#                   '% of Runs Working 2+ Routes']
#    else:
#        return None
#    return tr_keys, tr_type
#
#
## Applies Table Cell style to all cells in table
#def apply_table_cell(doc, table):
#    for r in table.rows:
#        for c in r.cells:
#            for p in c.paragraphs:
#                p.style = doc.styles['Table Cell']
