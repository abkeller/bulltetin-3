from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import os
from pathlib import Path

from ..data.rch_data import rch_data
from ..utils import pct, strint
from ..utils.doc_utils import *

GRAY = '#E7E7E7'

class BulletinCover:
    def __init__(self, curr_booking, prev_booking, export_dir):

        # Create document and add styles defined in doc_utils
        self.doc = Document()
        edit_styles(self.doc)
        set_margins(self.doc, Inches(0.5))

        # Calculate run characteristic data from OIG export
        self.run_char = {'curr': rch_data(curr_booking),
                         'prev': rch_data(prev_booking)}

        # Get pick names, start dates, and seasons
        self.curr_pick = self.run_char['curr']['booking_data']
        self.prev_pick = self.run_char['prev']['booking_data']

        # Add tables to document
        self.tables = {'extras': self.run_change_table('wk')}

        doc_name = '_Templates/bulletin.docx'
        doc_path = os.path.join(export_dir, doc_name)
        self.doc.save(doc_path)

        # Write report to PDF
        word = win32com.client.Dispatch('Word.Application')
        try:
            docx = word.Documents.Open(doc_path)
            try:
                pick_name = self.curr_pick['desc']
                pdf_name = 'Run Characteristics - General Stats - {}.pdf'.format(pick_name)
                pdf_path = os.path.join(export_dir, pdf_name)
                docx.SaveAs(pdf_path, FileFormat=17)
            except:
                print('      Error printing PDF report.')
            docx.Close()
        except:
            print('      Error printing PDF report.')

        word.Quit()

    # ************************************************************************
    # Pages 1-3: Run Change Profiles (Wk/Sa/Su)
    def run_change_table(self, day_type):
        
        self.add_headers('run_change_table', day_type=day_type)
        
        sub_head = 'Composite Changes in Run Categories for {} schedules'.format(DAY_TYPES[day_type]['name'])
        h2 = self.doc.add_heading(sub_head, level=2)
        h2.style = self.doc.styles['H2']
        
        curr_data = self.run_char['curr']['run_categories'][day_type]
        prev_data = self.run_char['prev']['run_categories'][day_type]

        # Create table and set headers
        if day_type == 'su':
            table = self.doc.add_table(rows=3 * len(curr_data) + 4, cols=4)
            set_col_widths(table, (Inches(1.5), Inches(1.5), Inches(1),
                                   Inches(1)))
        else:
            table = self.doc.add_table(rows=3 * len(curr_data) + 4, cols=5)
            set_col_widths(table, (Inches(1.5), Inches(1.5), Inches(1),
                                   Inches(1), Inches(1)))
        table.style = 'Table Grid'
        
       # create borders     
        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
        blank = {'val': 'none', 'color': '#ffffff'}
        row_length = len(table.rows)
        for i in range(row_length):
           for cell in table.rows[i].cells:
               set_cell_border(cell, start=thick_black, end=thick_black)
               
           if i == 0 or i == 2:
               for cell in table.rows[i].cells:
                   set_cell_border(cell, top=thick_black)
                   set_cell_border(table.columns[0].cells[0], start=blank, top=blank)
                   set_cell_border(table.columns[0].cells[1], start=blank, top=blank)

           if (i - 1) % 3 == 0:
               for cell in table.rows[i].cells:
                    set_cell_border(cell, bottom=thick_black)               
           if i > 22:
               for cell in table.rows[i].cells:
                   set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)
              
        # Merge first two rows except Saturday Tripper/PTO
        if day_type == 'wk':
            for c in range(5):
                header = table.cell(0, c).merge(table.cell(1, c))
                header.text = ('\nGarage\n', '\nPeriod\n', 'Regular Runs',
                               'Block', 'Total Trippers')[c]
            table.cell(0, 3).add_paragraph('Runs')
        elif day_type == 'sa':
            for c in range(3):
                header = table.cell(0, c).merge(table.cell(1, c))
                header.text = ('\nGarage\n', '\nPeriod\n', 'Regular Runs')[c]
            span = table.cell(0, 3).merge(table.cell(0, 4))
            span.text = 'Tripper/PTO Work'
            # add thick border under 'Tripper/PTO Work'
            set_cell_border(span, bottom=thick_black, top=thick_black)
            table.cell(1, 3).text = 'Trippers'
            table.cell(1, 4).text = 'Runs'
        else:
            for c in range(4):
                header = table.cell(0, c).merge(table.cell(1, c))
                header.text = ('\nGarage\n', '\nPeriod\n', 'Regular Runs',
                               'PTO')[c]
            table.cell(0, 3).add_paragraph('Work')

        row_names = [self.prev_pick['desc'].upper().split('BUS SYSTEM - ')[1], self.curr_pick['desc'].upper().split('BUS SYSTEM - ')[1], 'Difference']
        categories = {'wk': ['reg', 'block', 'tripper'],
                      'sa': ['reg', 'tripper', 'pto'],
                      'su': ['reg', 'pto']}[day_type]


        for i, g in enumerate(curr_data):
            # First column - Garage (merged)
            gar = table.cell(i * 3 + 2, 0).merge(table.cell(i * 3 + 4, 0))
            gar.text = GARAGES.get(g)

            # Second column - Period
            for j, r in enumerate(row_names):
                table.cell(i * 3 + j + 2, 1).text = r
                #name_cells = table.columns[1].cells((i * 2 + j + 2, 1), italic=True)

            # Third column - Regular Runs
            # Fourth column - Block Runs
            # Fifth column - Total Trippers
            for j, c in enumerate(categories):
                prev_ct, curr_ct = prev_data[g][c], curr_data[g][c]
                table.cell(i * 3 + 2, 2 + j).text = strint(prev_ct)
                table.cell(i * 3 + 3, 2 + j).text = strint(curr_ct)
                table.cell(i * 3 + 4, 2 + j).text = strint(curr_ct - prev_ct)

        # System totals
        sys_row = len(curr_data) * 3 + 2
        table.cell(sys_row, 0).text = 'SYSTEM TOTALS'
        table.cell(sys_row + 1, 0).text = 'Differentials'

        for j, c in enumerate(categories):
            prev_total = sum([prev_data[g][c] for g in prev_data])
            curr_total = sum([curr_data[g][c] for g in curr_data])
            table.cell(sys_row, 2 + j).text = strint(curr_total)
            table.cell(sys_row + 1, 2 + j).text = strint(
                curr_total - prev_total)

        # Format cells
        apply_table_cell(self.doc, table)

        # Bold and italic
        b_cells = table.columns[0].cells + table.columns[1].cells[3:-3:3]
        for r in [0, 1, -1, -2]:
            b_cells += table.rows[r].cells
        format_cells(b_cells, bold=True)

        i_cells = table.rows[-1].cells[2:] + table.columns[1].cells[2:]
        for r in table.rows[4:-2:3]:
            i_cells += r.cells[1:]
        format_cells(i_cells, italic=True)

        # Alignment
        for r in table.rows:
            format_cells(r.cells[2:], halign='right')
        if day_type == 'sa':
            format_cells(table.cell(0, 3), halign='center')

        # Gray shading
        g_rows = [table.rows[i:-2:6] for i in range(2, 5)]
        for r_list in g_rows:
            for r in r_list:
                format_cells(r.cells, shading=GRAY)

        return table

    # ************************************************************************
    # Pages 4-6: Tripper/Block/Swing Breakdown
    def breakdown_table(self, bd_type):
        self.add_headers(bd_type)
    
        data = self.run_char['curr'][bd_type]

        # Create table and set headers
        if bd_type == 'tripper_bd':
            row_count = 4 * (len(data) + 2) + 4
        else:
            row_count = 4 * (len(data) + 1) + 2
            
        table = self.doc.add_table(rows=row_count, cols=3)
        set_col_widths(table, (Inches(1.25), Inches(2.5), Inches(0.75)))
        table.style = 'Table Grid'

       # create borders     
        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
        blank = {'val': 'none', 'color': '#ffffff'}
        row_length = len(table.rows)
        for i in range(row_length):
           for cell in table.rows[i].cells:
               set_cell_border(cell, start=thick_black, end=thick_black)
               
               if i == 0:
                   set_cell_border(cell, top=thick_black)
                   set_cell_border(table.columns[0].cells[0], start=blank, top=blank)
                   
               if i > 28:
                   set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)


        col1 = 'Tripper Type\n' if bd_type == 'tripper_bd\n' else 'Run Type\n'
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = ('Garage\n', col1, 'Total')[i]
            format_cells(cell, bold=True)

        tr_keys, tr_type = get_keys(bd_type)

        for i, g in enumerate(data):
            # First column - Garage (merged)
            gar = table.cell(i * 4 + 1, 0).merge(table.cell(i * 4 + 4, 0))
            gar.text = GARAGES.get(g)
            format_cells(gar, bold=True)

            # Second column - Tripper/Run Type
            # Third column - Total
            for j in range(4):
                table.cell(i * 4 + 1 + j, 1).text = tr_type[j]
                is_pct = tr_keys[j] == 'pct'
                table.cell(i * 4 + 1 + j, 2).text = strint(
                    data[g][tr_keys[j]], pct=is_pct)

        # System total
        sys_start = len(data) * 4 + 2
        s_hdr = table.cell(sys_start, 0).merge(table.cell(sys_start + 3, 0))
        s_hdr.text = 'System'
        format_cells(s_hdr, bold=True)

        sys_total = [sum([data[g][tr_keys[j]] for g in data]) for j in range(4)]

        for j in range(4):
            table.cell(sys_start + j, 1).text = tr_type[j]
            
            # Change percentages to systemwide
            if j == 3 and tr_keys[j] == 'pct':
                if bd_type == 'block_pct':
                    table.cell(sys_start + 3, 2).text = strint(pct(
                        sys_total[1], sys_total[2]), pct=True)
                else:
                    table.cell(sys_start + 3, 2).text = strint(pct(
                        sys_total[0] + sys_total[1], sys_total[2]), pct=True)
            else:
                table.cell(sys_start + j, 2).text = strint(sys_total[j])

        # Calculation (Tripper BD only)
        if bd_type == 'tripper_bd':
            calc_start = sys_start + 5
            c_hdr = table.cell(calc_start, 0).merge(table.cell(
                calc_start + 4, 0))
            c_hdr.text = 'CALCULATION:'
            format_cells(c_hdr, bold=True)

            tr_multi = [5, 1, 2, 2]
            calc_type, calc = tr_type, sys_total

            for j in range(4):
                if tr_multi[j] > 1:
                    calc_type[j] += ' * {}'.format(tr_multi[j])
                    calc[j] = sys_total[j] * tr_multi[j]

                table.cell(calc_start + j, 1).text = calc_type[j]
                table.cell(calc_start + j, 2).text = strint(calc[j])

            c_ftr = table.cell(calc_start + 4, 1)
            c_ftr.text = 'CALCULATION:'
            ftr_calc = table.cell(calc_start + 4, 2)
            ftr_calc.text = strint(sum(calc) / 5)
            for cell in [c_ftr, ftr_calc]:
                format_cells(cell, bold=True)
            
            

        # Apply Table Cell style
        apply_table_cell(self.doc, table)

        # Draw thick borders and shade every other garage
        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
        for j in range(len(data) + 1):
            for cell in table.rows[j * 4 + 1].cells:
                set_cell_border(cell, top=thick_black)
            if j % 2 == 1:
                for k in range(4):
                    for cell in table.rows[j * 4 - k].cells:
                        set_cell_shading(cell, GRAY)

        # Add Calculation notes for Tripper Breakdown - TO-DO
        if bd_type == 'tripper_bd':
            calc_notes = ['','','Calculation = Total tripper pieces per week divided by 5',
                          'Weekday trippers are multiplied by 5 because there are 5 weekdays per week',
                          'Saturday and Sunday runs set aside for PTOs are multiplied by 2 because there are 2 pieces per run']

            for cn in calc_notes:
                para = self.doc.add_paragraph(cn)
                para.style = self.doc.styles['Footnote']

        return table

    # ************************************************************************
    # Pages 7-9: Interline Current Pick (Wk/Sa/Su)
    def interline_table(self, day_type):
        self.add_headers('interline_table', day_type=day_type)
        curr_data = self.run_char['curr']['interline'][day_type]

        # Create table and set headers
        cols = 4 if day_type == 'wk' else 3
        table = self.doc.add_table(rows=4 * len(curr_data) + 6, cols=cols)
        if day_type == 'wk':
            set_col_widths(table, (11.16, Inches(2.1), Inches(0.75),
                                   Inches(0.75)))
        else:
            set_col_widths(table, (Inches(0.9), Inches(2.1), Inches(1.25)))
        table.style = 'Table Grid'

       # create borders     
        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
        blank = {'val': 'none', 'color': '#ffffff'}
        row_length = len(table.rows)
        for i in range(row_length):
           for cell in table.rows[i].cells:
               set_cell_border(cell, start=thick_black, end=thick_black)
               
           if i == 0:
               for cell in table.rows[i].cells:
                   set_cell_border(cell, top=thick_black, bottom=blank)
                   set_cell_border(table.columns[i].cells[i], start=blank, top=blank)
                   
           if i == 1:
               for cell in table.rows[i].cells:
                   set_cell_border(cell, top=blank, bottom=thick_black)
                   set_cell_border(table.columns[0].cells[i], start=blank)
                  
           if (i - 5) % 4 == 0 and i >= 5:
               for cell in table.rows[i].cells:
                   set_cell_border(cell, top=thick_black, bottom=thick_black)
                   set_cell_border(table.columns[0].cells[i], start=blank)
                   set_cell_border(table.columns[cols - 1].cells[i])

           if i == 28:
               for cell in table.rows[i].cells:
                   set_cell_border(cell, bottom=thick_black)
                    
           if i > 28:
               for cell in table.rows[i].cells:
                   set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)

        # add column headers to table            
        headers = ('\nGarage', '\nRun Type',
                   '\nRegular Runs', '\nBlock')        
        for i, c in enumerate(headers):
            if day_type in ['su', 'sa'] and i == 4:
                pass
            elif day_type == 'su' and i == 2:
                c = '\nReg. FTO'
                table.cell(0, i).text = c
                format_cells(table.cell(0, i), bold=True)
            else:
                table.cell(0, i).text = c
                format_cells(table.cell(0, i), bold=True)

        tr_keys, tr_type = get_keys('interline')
        run_types = ['reg', 'block'] if day_type == 'wk' else ['reg']
        
        for i, g in enumerate(curr_data):
            # First column - Garage (merged)
            gar = table.cell(i * 4 + 2, 0).merge(table.cell(i * 4 + 4, 0))
            gar.text = GARAGES.get(g)
            format_cells(gar, bold=True)

            for j in range(3):
                # Second column - Run Type
                table.cell(i * 4 + 2 + j, 1).text = tr_type[j]

                # Third-Fifth columns - Regular Runs
                # Sixth-Eighth columns (Weekday) - Block Runs
                is_pct = tr_keys[j] == 'pct'
                row = i * 4 + 2 + j
                
                for k, run_type in enumerate(run_types):
                    curr_ct = curr_data[g][run_type][tr_keys[j]]
                    table.cell(row, 2 + k).text = strint(curr_ct, pct=is_pct)

        # System Total
        sys_start = len(curr_data) * 4 + 3
        system = table.cell(sys_start, 0).merge(table.cell(sys_start + 2, 0))
        system.text = 'System'
        format_cells(system, bold=True)

        for j in range(3):
            table.cell(sys_start + j, 1).text = tr_type[j]

        for i, run_type in enumerate(run_types):
            iline, total = [], []
            for j, data in enumerate(curr_data):
                #print(curr_data[data])
                iline.append(int(curr_data[data][run_type]['iline']))
                total.append(int(curr_data[data][run_type]['total']))
            iline_total = sum(iline)
            total = sum(total)

            col = i + 2
            table.cell(sys_start, col).text = strint(iline_total)
            table.cell(sys_start + 1, col).text = strint(total)
            table.cell(sys_start + 2, col).text = strint(pct(iline_total, total), pct=True)

        # Apply Table Cell style
        apply_table_cell(self.doc, table)

        # Draw thick borders and shade every other garage
        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
        for j in range(len(curr_data) + 1):
            for cell in table.rows[j * 3 + 2].cells:
                pass
                #set_cell_border(cell, top=thick_black)
            if j % 2 == 1:
                for k in range(3):
                    for cell in table.rows[j * 4 - k].cells:
                        set_cell_shading(cell, GRAY)

        return table


    # ************************************************************************
    # Pages 10-12: Interline Comparisons (Wk/Sa/Su)
    def interline_table_comp(self, day_type):
        self.add_headers('interline_table', day_type=day_type)
        curr_data = self.run_char['curr']['interline'][day_type]
        prev_data = self.run_char['prev']['interline'][day_type]

        # Create table and set headers
        cols = 8 if day_type == 'wk' else 5
        table = self.doc.add_table(rows=4 * len(curr_data) + 7, cols=cols)             
        
        if day_type == 'wk':
            set_col_widths(table, (Inches(0.9), Inches(2.1), Inches(0.75),
                                   Inches(0.75), Inches(0.75), Inches(0.75),
                                   Inches(0.75), Inches(0.75)))
        else:
            set_col_widths(table, (Inches(0.9), Inches(2.1), Inches(1.25),
                                   Inches(1.25), Inches(1.25)))
        table.style = 'Table Grid'
        row_length = len(table.rows)
       # create borders     
        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
        blank = {'val': 'none', 'color': '#ffffff'}
                 
        if day_type == 'wk':
            for i in range(row_length):
                if  i == 0:
                    for cell in table.rows[i].cells:
                        set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)
                        set_cell_border(table.columns[i].cells[i], start=blank, bottom=blank)
                
                if i == 1:
                    for cell in table.rows[i].cells:
                        set_cell_border(table.columns[0].cells[i], start=blank, bottom=blank, top=blank)

                if (i - 5) % 4 == 0 and i >= 5:
                   for cell in table.rows[i].cells:
                       #set_cell_border(cell, top=thick_black, bottom=thick_black)
                       set_cell_border(table.columns[0].cells[i], start=blank)
                       set_cell_border(table.columns[cols - 1].cells[i], end=blank)            

                if i > 28:
                   for cell in table.rows[i].cells:
                       set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)

        if day_type != 'wk':
            for i in range(row_length):
               for cell in table.rows[i].cells:
                   set_cell_border(cell, start=thick_black, end=thick_black)
                   
               if i == 0:
                   for cell in table.rows[i].cells:
                       set_cell_border(cell, top=blank, bottom=thick_black, start=blank, end=blank)
                       set_cell_border(table.columns[0].cells[i], start=blank, bottom=blank, top=blank)

                       
               if i == 1:
                   for cell in table.rows[i].cells:
                       set_cell_border(cell, top=blank, bottom=thick_black)
                      
               if (i - 5) % 4 == 0 and i >= 5:
                   for cell in table.rows[i].cells:
                       set_cell_border(cell, top=thick_black, bottom=thick_black)
                       set_cell_border(table.columns[0].cells[i], start=blank)
                       set_cell_border(table.columns[cols - 1].cells[i], end=blank)
                       set_cell_border(table.columns[0].cells[1], start=blank, top=blank)

    
               if i == 28:
                   for cell in table.rows[i].cells:
                       set_cell_border(cell, bottom=thick_black)
                        
               if i > 28:
                   for cell in table.rows[i].cells:
                       set_cell_border(cell, top=blank, bottom=blank, start=blank, end=blank)

        reg = table.cell(0, 2)
        reg.text = '{}\n'.format(self.prev_pick['desc'].upper().split('BUS SYSTEM - ')[1])
        curr = table.cell(0, 3)
        curr.text = '{}\n'.format(self.curr_pick['desc'].upper().split('BUS SYSTEM - ')[1])
        for cell in [reg, curr]:
            format_cells(cell, bold=True, italic=True)
        
        if day_type == 'wk':
            prev = table.cell(0, 5)
            prev.text = '{}\n'.format(self.prev_pick['desc'].upper().split('BUS SYSTEM - ')[1])
            curr = table.cell(0, 6)
            curr.text = '{}\n'.format(self.curr_pick['desc'].upper().split('BUS SYSTEM - ')[1]) 
            
            for cell in [prev, curr]:
                format_cells(cell, bold=True, italic=True)
            
        for i, c in enumerate(table.rows[1].cells):
            headers = ('Garage', 'Run Type\n\n', 'Regular Runs','Regular Runs', 'Change',
                       'Block', 'Block', 'Change')

            c.text = headers[i]
            format_cells(c, bold=True)

        tr_keys, tr_type = get_keys('interline')
        run_types = ['reg', 'block'] if day_type == 'wk' else ['reg']

        for i, g in enumerate(curr_data):
            # First column - Garage (merged)
            gar = table.cell(i * 4 + 2, 0).merge(table.cell(i * 4 + 4, 0))
            gar.text = GARAGES.get(g)
            format_cells(gar, bold=True)

            for j in range(3):
                # Second column - Run Type
                table.cell(i * 4 + 2 + j, 1).text = tr_type[j]

                # Third-Fifth columns - Regular Runs
                # Sixth-Eighth columns (Weekday) - Block Runs
                is_pct = tr_keys[j] == 'pct'
                row = i * 4 + 2 + j
                
                for k, run_type in enumerate(run_types):
                    prev_ct = prev_data[g][run_type][tr_keys[j]]
                    curr_ct = curr_data[g][run_type][tr_keys[j]]
                    table.cell(row, 2 + k * 3).text = strint(prev_ct, pct=is_pct)
                    table.cell(row, 3 + k * 3).text = strint(curr_ct, pct=is_pct)
                    table.cell(row, 4 + k * 3).text = strint(curr_ct - prev_ct,
                                                             pct=is_pct)

        # System Total
        sys_start = len(curr_data) * 4 + 4
        system = table.cell(sys_start, 0).merge(table.cell(sys_start + 2, 0))
        system.text = 'System'
        format_cells(system, bold=True)


        for j in range(3):
            table.cell(sys_start + j, 1).text = tr_type[j]

        for i, run_type in enumerate(run_types):
            iline, total = [], []
            for j, data in enumerate((prev_data, curr_data)):
                iline.append(sum([data[g][run_type]['iline'] for g in data]))
                total.append(sum([data[g][run_type]['total'] for g in data]))

                col = 2 + j + 3 * i
                table.cell(sys_start, col).text = strint(iline[j])
                table.cell(sys_start + 1, col).text = strint(total[j])
                table.cell(sys_start + 2, col).text = strint(pct(iline[j], total[j]), pct=True)

            col = 4 + 3 * i
            table.cell(sys_start, col).text = strint(iline[1] - iline[0])
            table.cell(sys_start + 1, col).text = strint(total[1] - total[0])
            table.cell(sys_start + 2, col).text = strint(pct(iline[1], total[1]) -
                                                         pct(iline[0], total[0]),
                                                         pct=True)

        # Apply Table Cell style
        apply_table_cell(self.doc, table)

        # Draw thick borders and shade every other garage
        thick_black = {'sz': 12, 'val': 'single', 'color': '#000000'}
        for j in range(len(curr_data) + 1):
            for cell in table.rows[j * 4 + 2].cells:
                pass
                #set_cell_border(cell, top=thick_black)
            if j % 2 == 1:
                for k in range(3):
                    for cell in table.rows[j * 4 - k].cells:
                        set_cell_shading(cell, GRAY)

        return table
    # ************************************************************************

    def add_headers(self, table_type, day_type=None):
        start_date = self.curr_pick['start_date'].strftime('%#d %B %Y').upper()
        h1 = self.doc.add_heading('CTA Bus System Effective {} ({})'.format(
            start_date, self.curr_pick['name']), level=1)
        h1.style = self.doc.styles['H1']

        if table_type == 'run_change_table':
            h2_text = 'Profile of Run Changes by Garage - {}'.format(
                DAY_TYPES[day_type]['name'].upper())
            h3_text = 'Composite Changes in Run Categories for {} Schedules'.format(
                DAY_TYPES[day_type]['name'])
        elif table_type == 'tripper_bd':
            h2_text = 'Breakdown of trippers and work set aside for part-time operators'
        elif table_type == 'block_pct':
            h2_text = 'Breakdown of Block Runs'
        elif table_type == 'swing_pct':
            h2_text = 'Breakdown of Swing Runs'
        elif table_type == 'interline_table':
            h2_text = 'Breakdown of Runs Working 2 or More Routes'        
        elif table_type == 'interline_table_comp':
            h2_text = 'Changes in Runs Working 2 or More Routes - {}'.format(
                DAY_TYPES[day_type]['name'].upper())
        else:
            h2_text = ''
        
        h2 = self.doc.add_heading(h2_text, level=2)
        h2.style = self.doc.styles['H2']

        if table_type == 'tripper_bd':
            h3_text = 'Calculated total of trippers must not exceed 1500'
            h3 = self.doc.add_paragraph(h3_text)
            h3.style = self.doc.styles['H2']
            h3.runs[0].font.italic = True

        elif table_type == 'block_pct':
            h3_text = 'Block Runs Must not exceed 20% of the total WK & Sa FTO Runs'
            h3 = self.doc.add_paragraph(h3_text)
            h3.style = self.doc.styles['H2']
            h3.runs[0].font.italic = True

        elif table_type == 'swing_pct':
            h3_text = 'Swing Runs must not exceed 30% of the total Wk & Sat FTO runs'
            h3 = self.doc.add_paragraph(h3_text)
            h3.style = self.doc.styles['H2']
            h3.runs[0].font.italic = True 
            
        elif table_type == 'interline_table':
            h3_text = DAY_TYPES[day_type]['name']
            h3 = self.doc.add_paragraph(h3_text)
            h3.style = self.doc.styles['H2']
            h3.runs[0].font.italic = True  

# Gets keys used to identify table data and their printed names
def get_keys(bd_type):
    if bd_type == 'tripper_bd':
        tr_keys = ['wk_tripper', 'sa_tripper', 'sa_pto', 'su_pto']
        tr_type = ['Weekday trippers',
                   'Saturday trippers',
                   'Saturday runs set aside for PTOs',
                   'Sunday runs set aside for PTOs']
    elif bd_type == 'block_pct':
        tr_keys = ['reg', 'block', 'fto', 'pct']
        tr_type = ['Non-block FTO Runs ((Wk*5)+Sa)',
                   'Block FTO Runs ((Wk*5)+Sa)',
                   'TOTAL FTO Runs ((Wk*5)+Sa)',
                   'Blocks as a % of Total']
    elif bd_type == 'swing_pct':
        tr_keys = ['sw', 'swb', 'fto', 'pct']
        tr_type = ['Regular Swing ((Wk*5)+Sa)',
                   'Block Swing ((Wk*5+Sa)',
                   'TOTAL FTO Runs ((Wk*5)+Sa)',
                   'Swings as a % of Total']
    elif bd_type == 'interline':
        tr_keys = ['iline', 'total', 'pct']
        tr_type = ['Runs working 2+ routes',
                   'TOTAL Runs',
                   '% of Runs Working 2+ Routes']
    else:
        return None
    return tr_keys, tr_type


# Applies Table Cell style to all cells in table
def apply_table_cell(doc, table):
    for r in table.rows:
        for c in r.cells:
            for p in c.paragraphs:
                p.style = doc.styles['Table Cell']
